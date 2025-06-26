import os
import re
import json
from datetime import datetime
import tempfile
from typing import List, Dict, Tuple
import io
import difflib

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR
from dotenv import load_dotenv

try:
    from PyPDF2 import PdfReader
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

# Load environment variables
load_dotenv()

# Configure Gemini AI
GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')

if not GEMINI_API_KEY:
    print("Warning: GEMINI_API_KEY not found in environment variables. Please add it to your .env file.")

try:
    import google.generativeai as genai
    genai.configure(api_key=GEMINI_API_KEY)
    model = genai.GenerativeModel('gemini-2.0-flash-exp')
except Exception as e:
    print(f"Error configuring Gemini AI: {e}")
    model = None

# Essay type classifications
ALLOWED_ESSAY_TYPES = [
    "Narrative Essay",
    "Descriptive Essay", 
    "Expository Essay",
    "Argumentative Essay",
    "Persuasive Essay",
    "Analytical Essay"
]

ESSAY_TYPES = {
    'Argumentative Essay': {
        'keywords': ['argue', 'thesis', 'evidence', 'counterargument', 'claim', 'support', 'oppose', 'debate'],
        'criteria': [
            'Clear thesis statement',
            'Strong evidence and examples',
            'Counterargument acknowledgment',
            'Logical flow of arguments',
            'Source credibility check'
        ]
    },
    'Narrative Essay': {
        'keywords': ['story', 'experience', 'happened', 'remember', 'narrative', 'personal', 'journey'],
        'criteria': [
            'Clear narrative arc',
            'Vivid imagery and descriptions',
            'Dialogue quality',
            'Character development',
            'Chronological flow'
        ]
    },
    'Analytical Essay': {
        'keywords': ['analyze', 'literary', 'author', 'character', 'theme', 'symbolism', 'literary device'],
        'criteria': [
            'Present tense usage',
            'Proper title italicization',
            'Quote integration',
            'Literary device identification',
            'Theme analysis depth'
        ]
    },
    'Expository Essay': {
        'keywords': ['explain', 'inform', 'describe', 'process', 'how to', 'definition'],
        'criteria': [
            'Clear explanations',
            'Logical organization',
            'Supporting details',
            'Objective tone',
            'Factual accuracy'
        ]
    },
    'Descriptive Essay': {
        'keywords': ['describe', 'imagery', 'sensory', 'vivid', 'details', 'scene', 'depict'],
        'criteria': [
            'Vivid sensory details',
            'Descriptive language',
            'Clear imagery',
            'Emotional impact',
            'Cohesive description'
        ]
    },
    'Persuasive Essay': {
        'keywords': ['persuade', 'convince', 'argument', 'position', 'appeal', 'rhetoric', 'call to action'],
        'criteria': [
            'Clear position statement',
            'Persuasive techniques (ethos, pathos, logos)',
            'Strong evidence',
            'Call to action',
            'Audience engagement'
        ]
    }
}

class DocumentProcessor:
    @staticmethod
    def read_docx(file_path):
        try:
            doc = Document(file_path)
            paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            return {
                'text': '\n\n'.join(paragraphs),
                'paragraphs': paragraphs,
                'paragraph_count': len(paragraphs)
            }
        except Exception as e:
            raise Exception(f"Error reading DOCX file: {str(e)}")
    
    @staticmethod
    def read_txt(file_path):
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
                return {
                    'text': content,
                    'paragraphs': paragraphs,
                    'paragraph_count': len(paragraphs)
                }
        except Exception as e:
            raise Exception(f"Error reading TXT file: {str(e)}")

    @staticmethod
    def read_pdf(file):
        if not PYPDF2_AVAILABLE:
            raise Exception("PyPDF2 is not installed. Please install it to process PDF files.")
        try:
            pdf = PdfReader(file)
            all_text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
            paragraphs = [p.strip() for p in all_text.split('\n\n') if p.strip()]
            return {
                'text': all_text,
                'paragraphs': paragraphs,
                'paragraph_count': len(paragraphs)
            }
        except Exception as e:
            raise Exception(f"Error reading PDF file: {str(e)}")

    @staticmethod
    def extract_text_from_docx(file):
        try:
            doc = Document(file)
            return "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")

    @staticmethod
    def extract_lines(file):
        try:
            if file.filename.lower().endswith('.docx'):
                return DocumentProcessor.extract_text_from_docx(file).splitlines()
            elif file.filename.lower().endswith('.txt'):
                return file.read().decode('utf-8').splitlines()
            elif file.filename.lower().endswith('.pdf'):
                if not PYPDF2_AVAILABLE:
                    raise Exception("PyPDF2 is not installed.")
                return DocumentProcessor.read_pdf(file)['text'].splitlines()
            else:
                return []
        except Exception as e:
            raise Exception(f"Error extracting lines: {str(e)}")

class EssayAnalyzer:
    def __init__(self):
        self.model = model
        self.essay_text = None
    
    def sanitize_text(self, text: str) -> str:
        """Sanitize text to ensure it can be safely included in JSON."""
        if not isinstance(text, str):
            return str(text)
        # Escape quotes, newlines, and other problematic characters
        text = text.replace('"', '\\"').replace('\n', '\\n').replace('\r', '\\r').replace('\t', '\\t')
        return text
    
    def extract_json(self, response_text: str) -> dict:
        """Extract and parse JSON from AI response, handling malformed JSON."""
        try:
            json_start = response_text.find('{')
            json_end = response_text.rfind('}') + 1
            if json_start == -1 or json_end == 0:
                print("Error: No valid JSON found in response")
                return None
            
            json_str = response_text[json_start:json_end].strip()
            json_str = re.sub(r'^```json\n|```$', '', json_str, flags=re.MULTILINE)
            # Parse JSON and sanitize string values
            parsed_json = json.loads(json_str)
            for key, value in parsed_json.items():
                if isinstance(value, str):
                    parsed_json[key] = self.sanitize_text(value)
                elif isinstance(value, list):
                    parsed_json[key] = [self.sanitize_text(item) if isinstance(item, str) else item for item in value]
            return parsed_json
        except json.JSONDecodeError as e:
            print(f"Error parsing JSON: {e}")
            return None
        except Exception as e:
            print(f"Unexpected error in extract_json: {e}")
            return None
    
    def classify_essay_type(self, text):
        text_lower = text.lower()
        scores = {}
        
        for essay_type, data in ESSAY_TYPES.items():
            score = sum(1 for keyword in data['keywords'] if keyword in text_lower)
            scores[essay_type] = score
        
        primary_type = max(scores, key=scores.get) if scores else 'Expository Essay'
        return primary_type
    
    def analyze_grammar_and_style(self, text):
        if not self.model:
            return self._fallback_analysis(text)
        
        prompt = f"""
        Analyze the following essay for grammar and style issues.
        Provide an overall score out of 100 and specific suggestions for improvement in JSON format:
        {{
            "overall_score": "score out of 100",
            "suggestions": ["suggestion 1", "suggestion 2", ...]
        }}
        
        Essay text:
        {self.sanitize_text(text[:2000])}...
        """
        
        try:
            response = self.model.generate_content(prompt)
            response_text = response.text
            json_start = response_text.find('{')
            json_end = response_text.rfind('}') + 1
            if json_start != -1 and json_end != -1:
                json_str = response_text[json_start:json_end]
                return json.loads(json_str)
            return self._fallback_analysis(text)
        except Exception as e:
            print(f"Error in analysis: {e}")
            return self._fallback_analysis(text)
    
    def create_correction_prompt(self, essay_text, essay_type):
        return f"""
        You are an expert writing assistant specializing in {essay_type}s.
        
        Analyze the following essay and provide corrections with tracked changes:
        1. Use <del>text</del> for text that should be deleted
        2. Use <ins>text</ins> for text that should be added
        3. Focus on grammar, style, clarity, and {essay_type.lower()} specific improvements
        4. Maintain the original meaning and structure
        5. Provide specific suggestions for improvement
        
        Return a valid JSON object:
        {{
          "essay_type": "{essay_type}",
          "essay_score": "numerical score out of 100",
          "corrected_essay": "Essay text with <del> and <ins> tags for tracked changes",
          "suggestions": ["specific suggestion 1", "specific suggestion 2", "specific suggestion 3"]
        }}
        
        Essay to analyze:
        \"\"\"
        {essay_text}
        \"\"\"
        """
    
    def _fallback_analysis(self, text):
        return {
            "overall_score": "75",
            "suggestions": ["Unable to perform detailed analysis due to missing AI model."]
        }

# Initialize Flask app and components
app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})
analyzer = EssayAnalyzer()
doc_processor = DocumentProcessor()

@app.route('/analyze_essay', methods=['POST'])
def analyze_essay():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        if not file.filename.lower().endswith('.docx'):
            return jsonify({'error': 'Only .docx files are supported'}), 400
        
        essay_text = doc_processor.extract_text_from_docx(file)
        if not essay_text or not essay_text.strip():
            return jsonify({'error': 'No text found in the document'}), 400
        if len(essay_text.strip()) < 50:
            return jsonify({'error': 'Essay is too short for meaningful analysis'}), 400

        if not analyzer.model:
            return jsonify({
                'essay_score': '75',
                'essay_type': 'Expository Essay',
                'corrected_essay': essay_text,
                'suggestions': ['AI model unavailable - please check your API key'],
            }), 503

        analyzer.essay_text = essay_text
        detected_type = analyzer.classify_essay_type(essay_text)
        
        # Get AI corrections and suggestions
        correction_prompt = analyzer.create_correction_prompt(essay_text, detected_type)
        correction_response = analyzer.model.generate_content(correction_prompt)
        corrected_json = analyzer.extract_json(correction_response.text)
        
        if not corrected_json:
            return jsonify({'error': 'Essay analysis failed - invalid AI response'}), 500
        
        return jsonify({
            'essay_score': corrected_json.get('essay_score', '75'),
            'essay_type': corrected_json.get('essay_type', detected_type),
            'corrected_essay': corrected_json.get('corrected_essay', essay_text),
            'suggestions': corrected_json.get('suggestions', [])
        }), 200

    except Exception as e:
        print(f"Error in essay analysis: {e}")
        return jsonify({'error': f'Analysis failed: {str(e)}'}), 500

@app.route('/change_essay_type', methods=['POST'])
def change_essay_type():
    try:
        data = request.get_json()
        essay_text = data.get('essay_text')
        target_type = data.get('target_essay_type')
        
        if not essay_text or not target_type:
            return jsonify({'error': 'essay_text and target_essay_type are required'}), 400
        if target_type not in ALLOWED_ESSAY_TYPES:
            return jsonify({'error': f'Invalid essay type. Allowed types: {ALLOWED_ESSAY_TYPES}'}), 400

        if not analyzer.model:
            return jsonify({
                'essay_score': '75',
                'essay_type': target_type,
                'corrected_essay': essay_text,
                'suggestions': ['AI model unavailable - please check your API key']
            }), 503

        # Clean the text of existing HTML tags
        clean_text = re.sub(r'<[^>]+>', '', essay_text)
        
        correction_prompt = analyzer.create_correction_prompt(clean_text, target_type)
        correction_response = analyzer.model.generate_content(correction_prompt)
        corrected_json = analyzer.extract_json(correction_response.text)
        
        if not corrected_json:
            return jsonify({'error': 'Essay type conversion failed - invalid AI response'}), 500

        return jsonify({
            'essay_score': corrected_json.get('essay_score', '75'),
            'essay_type': target_type,
            'corrected_essay': corrected_json.get('corrected_essay', essay_text),
            'suggestions': corrected_json.get('suggestions', [])
        }), 200

    except Exception as e:
        print(f"Error in essay type change: {e}")
        return jsonify({'error': f'Type conversion failed: {str(e)}'}), 500

@app.route('/download_revision', methods=['POST'])
def download_revision():
    try:
        data = request.get_json()
        final_text = data.get('final_text', '')
        title = data.get('title', 'Revised Essay')
        
        if not final_text:
            return jsonify({'error': 'No text provided for download'}), 400
        
        # Clean HTML tags from the text
        clean_text = re.sub(r'<[^>]+>', '', final_text)
        
        # Create Word document
        doc = Document()
        doc.add_heading(title, 0)
        
        # Add paragraphs
        paragraphs = clean_text.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Save to BytesIO
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return send_file(
            doc_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"{title.replace(' ', '_')}_final.docx"
        )
        
    except Exception as e:
        print(f"Error creating download: {e}")
        return jsonify({'error': 'Failed to create download file'}), 500

class DocumentComparator:
    @staticmethod
    def tag_word_diff(line1: str, line2: str) -> Tuple[str, str]:
        diff = difflib.ndiff(line1.split(), line2.split())
        tagged_line1, tagged_line2 = [], []
        
        for word in diff:
            if word.startswith('- '):
                tagged_line1.append(f"<del>{word[2:]}</del>")
            elif word.startswith('+ '):
                tagged_line2.append(f"<ins>{word[2:]}</ins>")
            elif word.startswith('  '):
                word_clean = word[2:]
                tagged_line1.append(word_clean)
                tagged_line2.append(word_clean)
        return ' '.join(tagged_line1), ' '.join(tagged_line2)

comparator = DocumentComparator()

@app.route('/compare_essays', methods=['POST'])
def compare_essays():
    try:
        if 'essay1' not in request.files or 'essay2' not in request.files:
            return jsonify({'error': 'Please upload both essay files'}), 400

        file1 = request.files['essay1']
        file2 = request.files['essay2']

        if not (file1.filename.lower().endswith(('.txt', '.docx', '.pdf')) and 
                file2.filename.lower().endswith(('.txt', '.docx', '.pdf'))):
            return jsonify({'error': 'Both files must be TXT, DOCX, or PDF format'}), 400

        lines1 = doc_processor.extract_lines(file1)
        lines2 = doc_processor.extract_lines(file2)

        if not model:
            return jsonify({'error': 'AI analysis unavailable - please check your API key'}), 503

        # Generate analysis for each essay
        full_essay1 = "\n".join(lines1)
        full_essay2 = "\n".join(lines2)
        
        analysis_prompt1 = f"""
        Analyze this essay and provide a brief summary of its characteristics (tone, style, structure, main ideas) in 2-3 sentences:
        
        {full_essay1[:1000]}...
        """
        
        analysis_prompt2 = f"""
        Analyze this essay and provide a brief summary of its characteristics (tone, style, structure, main ideas) in 2-3 sentences:
        
        {full_essay2[:1000]}...
        """
        
        try:
            response1 = model.generate_content(analysis_prompt1)
            essay1_analysis = response1.text.strip()
        except Exception as e:
            essay1_analysis = f"Error analyzing first essay: {str(e)}"
            
        try:
            response2 = model.generate_content(analysis_prompt2)
            essay2_analysis = response2.text.strip()
        except Exception as e:
            essay2_analysis = f"Error analyzing second essay: {str(e)}"

        # Generate summary of key differences
        summary_prompt = f"""
        Compare these two essays and identify key differences in tone, structure, ideas, and writing style.
        Provide 3-5 bullet points highlighting the main differences:

        Essay 1: {full_essay1[:500]}...
        Essay 2: {full_essay2[:500]}...
        """

        try:
            summary_response = model.generate_content(summary_prompt)
            key_differences = summary_response.text.strip()
        except Exception as e:
            key_differences = f"Error generating comparison: {str(e)}"

        return jsonify({
            'status': 'success',
            'draft1_analysis': essay1_analysis,
            'draft2_analysis': essay2_analysis,
            'key_differences': key_differences,
            'file1_name': secure_filename(file1.filename),
            'file2_name': secure_filename(file2.filename),
            'timestamp': datetime.now().isoformat()
        }), 200

    except Exception as e:
        print(f"Error in essay comparison: {e}")
        return jsonify({'error': f'Essay comparison failed: {str(e)}'}), 500

if __name__ == '__main__':
    print("Starting Essay Analyzer Server...")
    print("Make sure to set your GEMINI_API_KEY in a .env file")
    app.run(debug=True, host='0.0.0.0', port=5000)
